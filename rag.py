"""
rag.py — RAG pipeline: multi-format text extraction, semantic search, confidence, LLM via Groq.
CHANGES:
  - Added extract_text_from_file() supporting PDF, DOCX, DOC, XLSX, XLS, CSV, TXT, PNG, JPG, JPEG
  - Each parser has try/except with user-friendly st messages — never crashes the batch
  - _chunk_text() helper shared across all formats
  - PDF extraction and OCR fallback unchanged
"""

import io
import time
from typing import List, Tuple

import numpy as np
import requests
import streamlit as st

try:
    from sentence_transformers import SentenceTransformer
    SEMANTIC_AVAILABLE = True
except ImportError:
    SEMANTIC_AVAILABLE = False

try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False


@st.cache_resource(show_spinner=False)
def load_semantic_model():
    if not SEMANTIC_AVAILABLE:
        return None
    try:
        return SentenceTransformer('all-MiniLM-L6-v2')
    except Exception as e:
        st.warning(f"Semantic model failed: {e}. Falling back to keyword search.")
        return None


# ─────────────────────────────────────────────────────────────
# INTERNAL: chunk a raw text string into overlapping windows
# ─────────────────────────────────────────────────────────────
def _chunk_text(raw_text: str, filename: str,
                chunk_size: int = 100, overlap: int = 20) -> Tuple[List[str], bool]:
    """Split raw text into overlapping word-window chunks. Returns (chunks, used_ocr=False)."""
    words = raw_text.split()
    if len(words) < 10:
        st.warning(
            f"'{filename}': No readable text was found in this file. "
            "Try another file, a clearer scan, or convert it to PDF."
        )
        return [], False
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i + chunk_size])
        if len(chunk.strip()) > 30:
            chunks.append(chunk)
    return chunks, False


# ─────────────────────────────────────────────────────────────
# PDF  (existing logic, unchanged)
# ─────────────────────────────────────────────────────────────
def extract_text_from_pdf(pdf_bytes: bytes, filename: str) -> Tuple[List[str], bool]:
    chunk_size, overlap = 100, 20
    if PYPDF2_AVAILABLE:
        try:
            reader   = PdfReader(io.BytesIO(pdf_bytes))
            raw_text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    raw_text += t + " "
            words = raw_text.split()
            if len(words) > 50:
                chunks = []
                for i in range(0, len(words), chunk_size - overlap):
                    chunk = ' '.join(words[i:i + chunk_size])
                    if len(chunk.strip()) > 30:
                        chunks.append(chunk)
                return chunks, False
        except Exception as e:
            print(f"[rag] PyPDF2 parse error for '{filename}': {e}")

    if OCR_AVAILABLE:
        try:
            st.info(f"'{filename}' is scanned — running OCR (~30s per page)...")
            images   = convert_from_bytes(pdf_bytes, dpi=200)
            raw_text = ""
            for img in images:
                raw_text += pytesseract.image_to_string(img, lang='eng') + " "
            words = raw_text.split()
            if len(words) > 20:
                chunks = []
                for i in range(0, len(words), chunk_size - overlap):
                    chunk = ' '.join(words[i:i + chunk_size])
                    if len(chunk.strip()) > 30:
                        chunks.append(chunk)
                return chunks, True
        except Exception as e:
            st.warning(f"'{filename}': OCR failed — {e}")
    else:
        st.warning(
            f"'{filename}' appears to be a scanned PDF. "
            "Install pytesseract + pdf2image for OCR support."
        )

    return [], False


# ─────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────
def _extract_docx(file_bytes: bytes, filename: str) -> Tuple[List[str], bool]:
    try:
        import docx as _docx  # python-docx
    except ImportError:
        st.error(
            f"'{filename}': python-docx is not installed on the server. "
            "Run: pip install python-docx"
        )
        return [], False

    try:
        doc      = _docx.Document(io.BytesIO(file_bytes))
        parts    = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        # Include table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        raw_text = "\n".join(parts)
        return _chunk_text(raw_text, filename)
    except Exception as e:
        print(f"[rag] DOCX parse error for '{filename}': {e}")
        st.error(
            f"'{filename}': This file appears corrupted or unreadable. "
            "Please re-save the file and try again."
        )
        return [], False


# ─────────────────────────────────────────────────────────────
# DOC  (old binary format — best-effort via antiword, else prompt to convert)
# ─────────────────────────────────────────────────────────────
def _extract_doc(file_bytes: bytes, filename: str) -> Tuple[List[str], bool]:
    # Try antiword (if installed on server)
    try:
        import subprocess
        result = subprocess.run(
            ["antiword", "-"],
            input=file_bytes, capture_output=True, timeout=15
        )
        if result.returncode == 0 and result.stdout:
            raw_text = result.stdout.decode("utf-8", errors="replace")
            if raw_text.strip():
                return _chunk_text(raw_text, filename)
    except Exception as e:
        print(f"[rag] antiword attempt failed for '{filename}': {e}")

    # antiword not available or failed — ask user to convert
    st.warning(
        f"'{filename}': DOC files could not be read reliably. "
        "Please convert this file to DOCX or PDF and upload again."
    )
    return [], False


# ─────────────────────────────────────────────────────────────
# XLSX / XLS
# ─────────────────────────────────────────────────────────────
def _extract_excel(file_bytes: bytes, filename: str) -> Tuple[List[str], bool]:
    try:
        import pandas as pd  # noqa
    except ImportError:
        st.error(
            f"'{filename}': pandas / openpyxl are not installed. "
            "Run: pip install pandas openpyxl"
        )
        return [], False

    try:
        xl     = pd.ExcelFile(io.BytesIO(file_bytes))
        parts  = []
        for sheet_name in xl.sheet_names:
            try:
                df = xl.parse(sheet_name)
                if df.empty:
                    continue
                parts.append(f"[Sheet: {sheet_name}]")
                parts.append(df.to_string(index=False))
            except Exception as sheet_err:
                print(f"[rag] Skipping sheet '{sheet_name}' in '{filename}': {sheet_err}")
                continue

        if not parts:
            st.warning(
                f"'{filename}': The spreadsheet was read, "
                "but no usable text/content was found."
            )
            return [], False

        raw_text = "\n\n".join(parts)
        return _chunk_text(raw_text, filename)
    except Exception as e:
        print(f"[rag] Excel parse error for '{filename}': {e}")
        st.error(
            f"'{filename}': This file appears corrupted or unreadable. "
            "Please re-save the file and try again."
        )
        return [], False


# ─────────────────────────────────────────────────────────────
# CSV
# ─────────────────────────────────────────────────────────────
def _extract_csv(file_bytes: bytes, filename: str) -> Tuple[List[str], bool]:
    try:
        import pandas as pd  # noqa
    except ImportError:
        st.error(
            f"'{filename}': pandas is not installed. "
            "Run: pip install pandas"
        )
        return [], False

    df = None
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            df = pd.read_csv(io.BytesIO(file_bytes), encoding=encoding)
            break
        except Exception:
            continue

    if df is None:
        st.error(
            f"'{filename}': Could not decode CSV. "
            "Try saving the file as UTF-8 and uploading again."
        )
        return [], False

    if df.empty:
        st.warning(
            f"'{filename}': The CSV file is empty or has no usable data."
        )
        return [], False

    try:
        raw_text = df.to_string(index=False)
        return _chunk_text(raw_text, filename)
    except Exception as e:
        print(f"[rag] CSV to_string error for '{filename}': {e}")
        st.error(
            f"'{filename}': Could not process CSV content. "
            "Please check the file format and try again."
        )
        return [], False


# ─────────────────────────────────────────────────────────────
# TXT
# ─────────────────────────────────────────────────────────────
def _extract_txt(file_bytes: bytes, filename: str) -> Tuple[List[str], bool]:
    raw_text = None
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            raw_text = file_bytes.decode(encoding)
            break
        except Exception:
            continue

    if raw_text is None:
        st.error(
            f"'{filename}': Could not decode this text file. "
            "Please save it as UTF-8 and try again."
        )
        return [], False

    if not raw_text.strip():
        st.warning(f"'{filename}': The text file appears to be empty.")
        return [], False

    return _chunk_text(raw_text, filename)


# ─────────────────────────────────────────────────────────────
# Images (PNG / JPG / JPEG) — OCR
# ─────────────────────────────────────────────────────────────
def _extract_image(file_bytes: bytes, filename: str) -> Tuple[List[str], bool]:
    if not OCR_AVAILABLE:
        st.error(
            f"'{filename}': Image/PDF OCR support is not installed on the server. "
            "Please upload a text-based PDF or DOCX, "
            "or install OCR dependencies: pip install pytesseract pillow pdf2image"
        )
        return [], False

    try:
        from PIL import Image
        img      = Image.open(io.BytesIO(file_bytes))
        raw_text = pytesseract.image_to_string(img, lang='eng')
    except Exception as e:
        print(f"[rag] Image OCR error for '{filename}': {e}")
        st.error(
            f"'{filename}': Image OCR failed. "
            "Please upload a clearer image or a text-based document."
        )
        return [], False

    if not raw_text.strip() or len(raw_text.split()) < 10:
        st.warning(
            f"'{filename}': OCR found very little text in this image. "
            "Try uploading a clearer scan, a scanned PDF, or a text-based document."
        )
        return [], False

    chunks, _ = _chunk_text(raw_text, filename)
    return chunks, True   # mark used_ocr=True for images


# ─────────────────────────────────────────────────────────────
# PUBLIC: unified entry point
# ─────────────────────────────────────────────────────────────
def extract_text_from_file(file_bytes: bytes, filename: str) -> Tuple[List[str], bool]:
    """
    Extract text chunks from any supported file type.

    Supported: PDF, DOCX, DOC, XLSX, XLS, CSV, TXT, PNG, JPG, JPEG.
    Returns (chunks, used_ocr).
    Shows user-friendly st messages on errors — never raises.
    """
    if not filename or "." not in filename:
        st.error(
            f"'{filename}': Cannot determine file type (no extension). "
            "Please rename the file with the correct extension and try again."
        )
        return [], False

    ext = filename.rsplit(".", 1)[-1].lower()

    try:
        if ext == "pdf":
            return extract_text_from_pdf(file_bytes, filename)
        elif ext == "docx":
            return _extract_docx(file_bytes, filename)
        elif ext == "doc":
            return _extract_doc(file_bytes, filename)
        elif ext in ("xlsx", "xls"):
            return _extract_excel(file_bytes, filename)
        elif ext == "csv":
            return _extract_csv(file_bytes, filename)
        elif ext == "txt":
            return _extract_txt(file_bytes, filename)
        elif ext in ("png", "jpg", "jpeg"):
            return _extract_image(file_bytes, filename)
        else:
            st.error(
                f"'{filename}': This file type (.{ext}) is not supported. "
                "Please upload PDF, DOCX, XLSX, CSV, TXT, PNG, JPG, or JPEG."
            )
            return [], False
    except Exception as e:
        # Catch-all — should not normally reach here given per-parser try/except
        print(f"[rag] Unexpected error in extract_text_from_file for '{filename}': {e}")
        st.error(
            f"'{filename}': An unexpected error occurred while reading this file. "
            "Please re-save the file and try again."
        )
        return [], False


# ─────────────────────────────────────────────────────────────
# SEARCH + CONFIDENCE + LLM
# ─────────────────────────────────────────────────────────────
def semantic_search(query: str, model, all_embeddings: np.ndarray,
                    chunks: List[str], n_results: int = 5) -> Tuple[List[str], List[float]]:
    query_embedding = model.encode([query], normalize_embeddings=True)
    scores          = (query_embedding @ all_embeddings.T).flatten()
    top_indices     = np.argsort(scores)[-n_results:][::-1]
    results, result_scores = [], []
    for idx in top_indices:
        if scores[idx] > 0.15:
            results.append(chunks[idx])
            result_scores.append(float(scores[idx]))
    return results, result_scores


def compute_confidence(scores: List[float]) -> float:
    if not scores:
        return 0.0
    top      = scores[0]
    avg_top3 = np.mean(scores[:3]) if len(scores) >= 3 else np.mean(scores)
    return round(min(float(0.6 * top + 0.4 * avg_top3), 1.0), 3)


def confidence_html(confidence: float) -> str:
    pct = int(confidence * 100)
    if pct >= 65:
        bar_class, label, color = "conf-high",   "High confidence",                          "#00c9a7"
    elif pct >= 35:
        bar_class, label, color = "conf-medium", "Medium confidence",                        "#f0a500"
    else:
        bar_class, label, color = "conf-low",    "Low confidence — answer may be incomplete", "#f05252"
    return f"""
    <div class="conf-wrap">
        <div class="conf-label">Answer Confidence</div>
        <div class="conf-bg"><div class="conf-fill {bar_class}" style="width:{pct}%"></div></div>
        <div class="conf-pct" style="color:{color}">{pct}% — {label}</div>
    </div>"""


def generate_answer(query: str, context: List[str], api_key: str,
                    memory_context: str = "", lang_instruction: str = "") -> Tuple[str, bool]:
    context_text = "\n\n---\n\n".join(context)
    prompt = f"""You are a formal, accurate College AI Assistant for SRM Institute of Science and Technology's Department of Computer Science.

STRICT RULES:
1. Answer ONLY using information explicitly stated in the DOCUMENT CONTEXT below.
2. If the answer is not clearly present, respond ONLY with: "This information is not available in the current documents. Please contact the department office directly."
3. Do NOT guess, infer, or recommend external websites.
4. Be concise and factual.
5. Use bullet points only when listing multiple items.
{lang_instruction}

{memory_context}

DOCUMENT CONTEXT:
{context_text}

STUDENT QUESTION: {query}

ANSWER:"""

    for attempt in range(3):
        try:
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={"model": "llama-3.3-70b-versatile",
                      "messages": [{"role": "user", "content": prompt}],
                      "temperature": 0.3, "max_tokens": 600},
                timeout=30
            )
            if response.status_code == 200:
                data = response.json()
                if data.get('choices') and data['choices'][0].get('message'):
                    return data['choices'][0]['message']['content'], True
                return "Empty response from AI model. Please try again.", False
            elif response.status_code == 401:
                return "Invalid Groq API key.", False
            elif response.status_code == 429:
                if attempt < 2:
                    time.sleep(2 ** attempt); continue
                return "Rate limit reached. Please wait a moment.", False
            elif response.status_code >= 500:
                if attempt < 2:
                    time.sleep(1); continue
                return f"Groq server error ({response.status_code}).", False
            else:
                return f"API error: HTTP {response.status_code}", False
        except requests.exceptions.Timeout:
            if attempt < 2:
                time.sleep(1); continue
            return "Request timed out.", False
        except requests.exceptions.ConnectionError:
            return "Cannot reach Groq API. Check your internet connection.", False
        except Exception as e:
            return f"Unexpected error: {e}", False
    return "All retry attempts failed. Please try again later.", False